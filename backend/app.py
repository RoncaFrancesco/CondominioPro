from flask import Flask, request, jsonify, send_from_directory, make_response
from flask_cors import CORS
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO

# Import moduli locali
from database_universal import init_db, create_default_user
from models import User, Condominio, Persona, Spesa, Millesemo, PreventivoAnnuale, SpesaPreventivata, RipartizionePreventivo, UnitaImmobiliare
from utils import (
    token_required, hash_password, generate_jwt_token, verify_jwt_token,
    validate_login_data, validate_condominio_data, validate_persona_data,
    validate_spesa_data, validate_millesimi_data, calculate_ripartizione_completa,
    calculate_ripartizione_preventivo, export_condominio_json, generate_preventivo_anno, log_error
)

# Inizializza Flask
app = Flask(__name__, static_folder='../frontend', static_url_path='')
CORS(app)

# Inizializza database all'avvio
with app.app_context():
    init_db()
    create_default_user()

# Servi file statici (frontend)
@app.route('/')
def index():
    return send_from_directory(app.static_folder, 'index.html')

# ======================
# AUTHENTICATION ENDPOINTS
# ======================

@app.route('/api/login', methods=['POST'])
def login():
    """Login utente"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        # Validazione dati
        errors = validate_login_data(data)
        if errors:
            return jsonify({'message': 'Dati non validi', 'errors': errors}), 400

        # Trova utente
        user = User.find_by_username(data['username'])
        if not user:
            return jsonify({'message': 'Credenziali non valide'}), 401

        # Verifica password (in produzione usare password hashing)
        if user.password != data['password']:
            return jsonify({'message': 'Credenziali non valide'}), 401

        # Genera token
        token = generate_jwt_token(user.id, user.username)

        return jsonify({
            'message': 'Login successful',
            'token': token,
            'user': {
                'id': user.id,
                'username': user.username
            }
        }), 200

    except Exception as e:
        log_error(str(e), 'login')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/logout', methods=['POST'])
@token_required
def logout():
    """Logout utente"""
    return jsonify({'message': 'Logout successful'}), 200

@app.route('/api/register', methods=['POST'])
def register():
    """Registrazione nuovo utente"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        username = data.get('username', '').strip()
        password = data.get('password', '').strip()

        # Validazioni
        if not username or len(username) < 3:
            return jsonify({'message': 'Username deve avere almeno 3 caratteri'}), 400
        if not password or len(password) < 6:
            return jsonify({'message': 'Password deve avere almeno 6 caratteri'}), 400

        # Controlla se username esiste già
        if User.find_by_username(username):
            return jsonify({'message': 'Username già in uso'}), 400

        # Crea nuovo utente
        user = User(username=username, password=password)  # In produzione: hash_password(password)
        user.save()

        return jsonify({'message': 'Utente creato con successo'}), 201

    except Exception as e:
        log_error(str(e), 'register')
        return jsonify({'message': 'Errore del server'}), 500

# ======================
# CONDOMINII ENDPOINTS
# ======================

@app.route('/api/condominii', methods=['GET'])
@token_required
def get_condominii():
    """Lista condominii utente"""
    try:
        user_id = request.current_user_id
        condominii = Condominio.get_by_user_id(user_id)

        result = []
        for condo in condominii:
            result.append({
                'id': condo.id,
                'nome': condo.nome,
                'indirizzo': condo.indirizzo,
                'num_unita': condo.num_unita,
                'anno_costruzione': condo.anno_costruzione,
                'numero_scale': condo.numero_scale,
                'presidente_assemblea': condo.presidente_assemblea,
                'responsabile': condo.responsabile,
                'telefono_responsabile': condo.telefono_responsabile,
                'email_responsabile': condo.email_responsabile,
                'amministratore_esterno': condo.amministratore_esterno,
                'iban_condominio': condo.iban_condominio,
                'banca_appoggio': condo.banca_appoggio,
                'created_at': condo.created_at
            })

        return jsonify(result), 200

    except Exception as e:
        log_error(str(e), 'get_condominii')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii', methods=['POST'])
@token_required
def create_condominio():
    """Crea nuovo condominio"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        # Validazione
        errors = validate_condominio_data(data)
        if errors:
            return jsonify({'message': 'Dati non validi', 'errors': errors}), 400

        # Crea condominio
        condominio = Condominio(
            user_id=request.current_user_id,
            nome=data['nome'],
            indirizzo=data.get('indirizzo', ''),
            num_unita=data['num_unita']
        )
        condominio.save()

        return jsonify({
            'message': 'Condominio creato con successo',
            'condominio': {
                'id': condominio.id,
                'nome': condominio.nome,
                'indirizzo': condominio.indirizzo,
                'num_unita': condominio.num_unita
            }
        }), 201

    except Exception as e:
        log_error(str(e), 'create_condominio')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>', methods=['GET'])
@token_required
def get_condominio(condo_id):
    """Dettagli condominio"""
    try:
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404

        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        return jsonify({
            'id': condominio.id,
            'nome': condominio.nome,
            'indirizzo': condominio.indirizzo,
            'num_unita': condominio.num_unita,
            'created_at': condominio.created_at
        }), 200

    except Exception as e:
        log_error(str(e), f'get_condominio {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>', methods=['PUT'])
@token_required
def update_condominio(condo_id):
    """Modifica condominio"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404

        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Aggiorna campi
        if 'nome' in data:
            condominio.nome = data['nome']
        if 'indirizzo' in data:
            condominio.indirizzo = data['indirizzo']
        if 'num_unita' in data:
            # Note: il cambio di num_unita richiede logica complessa
            # per ora non permettiamo questa modifica
            return jsonify({'message': 'Modifica numero unità non supportata'}), 400

        condominio.save()

        return jsonify({
            'message': 'Condominio aggiornato con successo',
            'condominio': {
                'id': condominio.id,
                'nome': condominio.nome,
                'indirizzo': condominio.indirizzo,
                'num_unita': condominio.num_unita
            }
        }), 200

    except Exception as e:
        log_error(str(e), f'update_condominio {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>', methods=['DELETE'])
@token_required
def delete_condominio(condo_id):
    """Elimina condominio"""
    try:
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404

        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        condominio.delete()

        return jsonify({'message': 'Condominio eliminato con successo'}), 200

    except Exception as e:
        log_error(str(e), f'delete_condominio {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

# ======================
# UNITA IMMOBILIARI ENDPOINTS
# ======================

@app.route('/api/condominii/<int:condo_id>/unita', methods=['GET'])
@token_required
def get_unita_condominio(condo_id):
    """Elenco unità immobiliari del condominio"""
    try:
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        unita = UnitaImmobiliare.get_by_condominio(condo_id)

        result = []
        for u in unita:
            result.append({
                'id': u.id,
                'numero_unita': u.numero_unita
            })

        return jsonify(result), 200

    except Exception as e:
        log_error(str(e), f'get_unita_condominio {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

# ======================
# PERSONE ENDPOINTS
# ======================

@app.route('/api/condominii/<int:condo_id>/persone', methods=['GET'])
@token_required
def get_persone(condo_id):
    """Lista persone condominio"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        persone = Persona.get_by_condominio(condo_id)

        result = []
        for persona in persone:
            result.append({
                'id': persona.id,
                'nome': persona.nome,
                'cognome': persona.cognome,
                'email': persona.email,
                'tipo_persona': persona.tipo_persona,
                'unita_id': persona.unita_id
            })

        return jsonify(result), 200

    except Exception as e:
        log_error(str(e), f'get_persone {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>/persone', methods=['POST'])
@token_required
def create_persona(condo_id):
    """Crea persona"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        # Validazione
        errors = validate_persona_data(data)
        if errors:
            return jsonify({'message': 'Dati non validi', 'errors': errors}), 400

        # Crea persona
        persona = Persona(
            condominio_id=condo_id,
            unita_id=data['unita_id'],
            nome=data['nome'],
            cognome=data['cognome'],
            email=data.get('email', ''),
            tipo_persona=data['tipo_persona']
        )
        persona.save()

        return jsonify({
            'message': 'Persona creata con successo',
            'persona': {
                'id': persona.id,
                'nome': persona.nome,
                'cognome': persona.cognome,
                'email': persona.email,
                'tipo_persona': persona.tipo_persona,
                'unita_id': persona.unita_id
            }
        }), 201

    except Exception as e:
        log_error(str(e), f'create_persona {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/persone/<int:persona_id>', methods=['PUT'])
@token_required
def update_persona(persona_id):
    """Modifica persona"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        persona = Persona.find_by_id(persona_id)
        if not persona:
            return jsonify({'message': 'Persona non trovata'}), 404

        # Verifica autorizzazione
        condominio = Condominio.find_by_id(persona.condominio_id)
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Aggiorna campi
        if 'unita_id' in data:
            persona.unita_id = data['unita_id']
        if 'nome' in data:
            persona.nome = data['nome']
        if 'cognome' in data:
            persona.cognome = data['cognome']
        if 'email' in data:
            persona.email = data['email']
        if 'tipo_persona' in data:
            persona.tipo_persona = data['tipo_persona']

        persona.save()

        return jsonify({
            'message': 'Persona aggiornata con successo',
            'persona': {
                'id': persona.id,
                'nome': persona.nome,
                'cognome': persona.cognome,
                'email': persona.email,
                'tipo_persona': persona.tipo_persona,
                'unita_id': persona.unita_id
            }
        }), 200

    except Exception as e:
        log_error(str(e), f'update_persona {persona_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/persone/<int:persona_id>', methods=['DELETE'])
@token_required
def delete_persona(persona_id):
    """Elimina persona"""
    try:
        persona = Persona.find_by_id(persona_id)
        if not persona:
            return jsonify({'message': 'Persona non trovata'}), 404

        # Verifica autorizzazione
        condominio = Condominio.find_by_id(persona.condominio_id)
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        persona.delete()

        return jsonify({'message': 'Persona eliminata con successo'}), 200

    except Exception as e:
        log_error(str(e), f'delete_persona {persona_id}')
        return jsonify({'message': 'Errore del server'}), 500

# ======================
# MILLESIMI ENDPOINTS
# ======================

@app.route('/api/condominii/<int:condo_id>/millesimi', methods=['GET'])
@token_required
def get_millesimi(condo_id):
    """Tutti millesimi condominio"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Ottieni millesimi per tutte le tabelle
        millesimi_completi = {}
        for tabella in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'L']:
            millesimi_completi[tabella] = Millesemo.get_by_condominio_tabella(condo_id, tabella)

        # Formatta risultato
        result = {}
        for tabella, millesimi_list in millesimi_completi.items():
            result[tabella] = []
            for millesimo in millesimi_list:
                result[tabella].append({
                    'id': millesimo.id,
                    'unita_id': millesimo.unita_id,
                    'valore': millesimo.valore
                })

        return jsonify(result), 200

    except Exception as e:
        log_error(str(e), f'get_millesimi {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>/millesimi/<tabella>', methods=['GET'])
@token_required
def get_millesimi_tabella(condo_id, tabella):
    """Millesimi per tabella specifica"""
    try:
        if tabella not in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'L']:
            return jsonify({'message': 'Tabella non valida'}), 400

        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        millesimi = Millesemo.get_by_condominio_tabella(condo_id, tabella)

        result = []
        for millesimo in millesimi:
            result.append({
                'id': millesimo.id,
                'unita_id': millesimo.unita_id,
                'valore': millesimo.valore
            })

        return jsonify(result), 200

    except Exception as e:
        log_error(str(e), f'get_millesimi_tabella {condo_id} {tabella}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>/millesimi', methods=['POST'])
@token_required
def save_millesimi(condo_id):
    """Crea/aggiorna millesimi"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        tabella = data.get('tabella')
        if not tabella or tabella not in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'L']:
            return jsonify({'message': 'Tabella non valida'}), 400

        millesimi_list = data.get('millesimi', [])
        errors = validate_millesimi_data(millesimi_list, condominio.num_unita)
        if errors:
            return jsonify({'message': 'Dati non validi', 'errors': errors}), 400

        # Salva millesimi
        for millesimo_data in millesimi_list:
            millesimo = Millesemo(
                condominio_id=condo_id,
                unita_id=millesimo_data['unita_id'],
                tabella=tabella,
                valore=millesimo_data['valore']
            )
            millesimo.save()

        return jsonify({'message': 'Millesimi salvati con successo'}), 201

    except Exception as e:
        log_error(str(e), f'save_millesimi {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>/millesimi/bulk', methods=['POST'])
@token_required
def save_millesimi_bulk(condo_id):
    """Salva millesimi con validazione e feedback per riga"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        tabella = data.get('tabella')
        millesimi_list = data.get('millesimi', [])

        if not tabella or tabella not in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'L']:
            return jsonify({'message': 'Tabella non valida'}), 400

        # Validazione complessiva
        errors = validate_millesimi_data(millesimi_list, condominio.num_unita)
        per_item = []

        # Validazione per riga
        for i, m in enumerate(millesimi_list):
            item_errors = []
            if not isinstance(m, dict):
                item_errors.append('Formato non valido')
            else:
                if 'unita_id' not in m:
                    item_errors.append('unita_id mancante')
                if 'valore' not in m:
                    item_errors.append('valore mancante')
                elif not isinstance(m['valore'], (int, float)) or m['valore'] < 0:
                    item_errors.append('valore non valido (>=0)')
            per_item.append({'index': i, 'errors': item_errors})

        if errors:
            return jsonify({
                'message': 'Dati non validi',
                'errors': errors,
                'per_item': per_item
            }), 400

        # Salvataggio
        for m in millesimi_list:
            millesimo = Millesemo(
                condominio_id=condo_id,
                unita_id=m['unita_id'],
                tabella=tabella,
                valore=m['valore']
            )
            millesimo.save()

        totale = sum(m['valore'] for m in millesimi_list if isinstance(m, dict) and 'valore' in m)

        return jsonify({
            'message': 'Millesimi salvati con successo',
            'totale': totale,
            'per_item': per_item
        }), 201

    except Exception as e:
        log_error(str(e), f'save_millesimi_bulk {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>/millesimi/validazione', methods=['GET'])
@token_required
def validate_millesimi_totali(condo_id):
    """Valida totale millesimi = 1000 per tutte le tabelle"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        result = {}
        for tabella in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'L']:
            result[tabella] = Millesemo.validate_total(condo_id, tabella)

        return jsonify({'validazione': result}), 200

    except Exception as e:
        log_error(str(e), f'validate_millesimi_totali {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

# ======================
# SPESE ENDPOINTS
# ======================

@app.route('/api/condominii/<int:condo_id>/spese', methods=['GET'])
@token_required
def get_spese(condo_id):
    """Lista spese condominio con filtro opzionale per tabella"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Filtro opzionale per tabella
        tabella_filter = request.args.get('tabella')
        if tabella_filter and tabella_filter not in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'L']:
            return jsonify({'message': 'Tabella non valida'}), 400

        spese = Spesa.get_by_condominio(condo_id, tabella_filter)

        result = []
        for spesa in spese:
            result.append({
                'id': spesa.id,
                'descrizione': spesa.descrizione,
                'importo': spesa.importo,
                'data_spesa': spesa.data_spesa,
                'tabella_millesimi': spesa.tabella_millesimi,
                'logica_pi': spesa.logica_pi,
                'percentuale_proprietario': spesa.percentuale_proprietario,
                'percentuale_inquilino': spesa.percentuale_inquilino,
                'created_at': spesa.created_at
            })

        return jsonify(result), 200

    except Exception as e:
        log_error(str(e), f'get_spese {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>/spese', methods=['POST'])
@token_required
def create_spesa(condo_id):
    """Crea spesa"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        # Validazione
        errors = validate_spesa_data(data)
        if errors:
            return jsonify({'message': 'Dati non validi', 'errors': errors}), 400

        # Crea spesa
        spesa = Spesa(
            condominio_id=condo_id,
            descrizione=data['descrizione'],
            importo=data['importo'],
            data_spesa=data.get('data_spesa', datetime.now().strftime('%Y-%m-%d')),
            tabella_millesimi=data['tabella_millesimi'],
            logica_pi=data['logica_pi'],
            percentuale_proprietario=data.get('percentuale_proprietario', 100),
            percentuale_inquilino=data.get('percentuale_inquilino', 0)
        )
        spesa.save()

        return jsonify({
            'message': 'Spesa creata con successo',
            'spesa': {
                'id': spesa.id,
                'descrizione': spesa.descrizione,
                'importo': spesa.importo,
                'data_spesa': spesa.data_spesa,
                'tabella_millesimi': spesa.tabella_millesimi,
                'logica_pi': spesa.logica_pi,
                'percentuale_proprietario': spesa.percentuale_proprietario,
                'percentuale_inquilino': spesa.percentuale_inquilino
            }
        }), 201

    except Exception as e:
        log_error(str(e), f'create_spesa {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/spese/<int:spesa_id>', methods=['PUT'])
@token_required
def update_spesa(spesa_id):
    """Modifica spesa"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        spesa = Spesa.find_by_id(spesa_id)
        if not spesa:
            return jsonify({'message': 'Spesa non trovata'}), 404

        # Verifica autorizzazione
        condominio = Condominio.find_by_id(spesa.condominio_id)
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Aggiorna campi
        if 'descrizione' in data:
            spesa.descrizione = data['descrizione']
        if 'importo' in data:
            spesa.importo = data['importo']
        if 'tabella_millesimi' in data:
            spesa.tabella_millesimi = data['tabella_millesimi']
        if 'logica_pi' in data:
            spesa.logica_pi = data['logica_pi']
        if 'percentuale_proprietario' in data:
            spesa.percentuale_proprietario = data['percentuale_proprietario']
        if 'percentuale_inquilino' in data:
            spesa.percentuale_inquilino = data['percentuale_inquilino']

        spesa.save()

        return jsonify({
            'message': 'Spesa aggiornata con successo',
            'spesa': {
                'id': spesa.id,
                'descrizione': spesa.descrizione,
                'importo': spesa.importo,
                'tabella_millesimi': spesa.tabella_millesimi,
                'logica_pi': spesa.logica_pi,
                'percentuale_proprietario': spesa.percentuale_proprietario,
                'percentuale_inquilino': spesa.percentuale_inquilino
            }
        }), 200

    except Exception as e:
        log_error(str(e), f'update_spesa {spesa_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/spese/<int:spesa_id>', methods=['DELETE'])
@token_required
def delete_spesa(spesa_id):
    """Elimina spesa"""
    try:
        spesa = Spesa.find_by_id(spesa_id)
        if not spesa:
            return jsonify({'message': 'Spesa non trovata'}), 404

        # Verifica autorizzazione
        condominio = Condominio.find_by_id(spesa.condominio_id)
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        spesa.delete()

        return jsonify({'message': 'Spesa eliminata con successo'}), 200

    except Exception as e:
        log_error(str(e), f'delete_spesa {spesa_id}')
        return jsonify({'message': 'Errore del server'}), 500

# ======================
# RIPARTIZIONE ENDPOINTS
# ======================

@app.route('/api/condominii/<int:condo_id>/ripartizione', methods=['GET'])
@token_required
def get_ripartizione(condo_id):
    """Calcola e ritorna ripartizione totale o per tabella specifica"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Filtro opzionale per tabella
        tabella_filter = request.args.get('tabella')
        if tabella_filter and tabella_filter not in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'L']:
            return jsonify({'message': 'Tabella non valida'}), 400

        # Calcola ripartizione completa o per tabella specifica
        if tabella_filter:
            # Calcola solo per la tabella specifica
            from database_universal import get_db
            conn = get_db()
            cursor = conn.cursor()

            # Assicura che la tabella ripartizione sia aggiornata
            try:
                calculate_ripartizione_completa(condo_id)
            except Exception as _:
                pass

            # Calcola ripartizione solo per quella tabella
            persone = Persona.get_by_condominio(condo_id)
            ripartizione = {}

            for persona in persone:
                # Filtra spese solo per quella tabella
                cursor.execute("""
                    SELECT s.*, rs.importo_dovuto
                    FROM spese s
                    LEFT JOIN ripartizione_spese rs ON s.id = rs.spesa_id
                    WHERE s.condominio_id = ? AND s.tabella_millesimi = ? AND rs.persona_id = ?
                    ORDER BY s.created_at
                """, (condo_id, tabella_filter, persona.id))

                totale_persona = 0
                for row in cursor.fetchall():
                    if row['importo_dovuto']:
                        totale_persona += row['importo_dovuto']

                if totale_persona > 0:
                    ripartizione[persona.id] = totale_persona

            conn.close()
        else:
            # Calcola ripartizione completa per tutte le tabelle
            ripartizione = calculate_ripartizione_completa(condo_id)

        # Formatta risultato con dettagli persone
        persone = Persona.get_by_condominio(condo_id)
        result = []

        for persona in persone:
            importo_dovuto = ripartizione.get(persona.id, 0)
            if importo_dovuto > 0:
                result.append({
                    'persona_id': persona.id,
                    'nome': persona.nome,
                    'cognome': persona.cognome,
                    'tipo_persona': persona.tipo_persona,
                    'importo_dovuto': importo_dovuto
                })

        return jsonify({
            'ripartizione': result,
            'totale': sum(ripartizione.values()),
            'tabella_filter': tabella_filter
        }), 200

    except Exception as e:
        log_error(str(e), f'get_ripartizione {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>/ripartizione/dettagliata', methods=['GET'])
@token_required
def get_ripartizione_dettagliata(condo_id):
    """Calcola e ritorna ripartizione dettagliata per persona con suddivisione per tabella"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Filtro opzionale per tabella
        tabella_filter = request.args.get('tabella')
        if tabella_filter and tabella_filter not in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'L']:
            return jsonify({'message': 'Tabella non valida'}), 400

        # Implementazione con dati reali dal database
        from database_universal import get_db
        conn = get_db()
        cursor = conn.cursor()

        # 1. Ottieni tutte le persone del condominio con numero unità
        cursor.execute('''
            SELECT p.*, ui.numero_unita
            FROM persone p
            JOIN unita_immobiliari ui ON p.unita_id = ui.id
            WHERE p.condominio_id = ?
            ORDER BY ui.numero_unita, p.cognome, p.nome
        ''', (condo_id,))

        persone_list = cursor.fetchall()

        # Mappa di ruoli presenti per unità per gestire il caso 50/50
        unita_ruoli = {}
        for p in persone_list:
            uid = p['unita_id']
            if uid not in unita_ruoli:
                unita_ruoli[uid] = set()
            if p['tipo_persona'] == 'proprietario_inquilino':
                unita_ruoli[uid].update({'proprietario', 'inquilino'})
            else:
                unita_ruoli[uid].add(p['tipo_persona'])

        # 2. Ottieni le spese del condominio (con eventuale filtro tabella)
        if tabella_filter:
            cursor.execute('''
                SELECT s.*, DATE(s.data_spesa) as data_formatted
                FROM spese s
                WHERE s.condominio_id = ? AND s.tabella_millesimi = ?
                ORDER BY s.data_spesa, s.descrizione
            ''', (condo_id, tabella_filter))
        else:
            cursor.execute('''
                SELECT s.*, DATE(s.data_spesa) as data_formatted
                FROM spese s
                WHERE s.condominio_id = ?
                ORDER BY s.data_spesa, s.descrizione
            ''', (condo_id,))

        spese_list = cursor.fetchall()

        result = []
        totale_generale = 0.0

        for persona in persone_list:
            persona_data = {
                'persona_id': persona['id'],
                'nome': persona['nome'],
                'cognome': persona['cognome'],
                'tipo_persona': persona['tipo_persona'],
                'numero_unita': persona['numero_unita'],
                'totale_dovuto': 0.0,
                'spese_per_tabella': {}
            }

            for spesa in spese_list:
                # Ottieni millesimi dell'unità per questa tabella
                cursor.execute('''
                    SELECT valore FROM millesimi
                    WHERE unita_id = ? AND tabella = ? AND condominio_id = ?
                ''', (persona['unita_id'], spesa['tabella_millesimi'], condo_id))

                millesimo_result = cursor.fetchone()
                if not millesimo_result:
                    continue

                millesimi = millesimo_result['valore']

                # Calcola percentuale in base alla logica P/I
                if spesa['logica_pi'] == 'proprietario':
                    percentuale = 100 if persona['tipo_persona'] in ['proprietario', 'proprietario_inquilino'] else 0
                elif spesa['logica_pi'] == 'inquilino':
                    percentuale = 100 if persona['tipo_persona'] in ['inquilino', 'proprietario_inquilino'] else 0
                elif spesa['logica_pi'] == '50/50':
                    # Se persona ricopre entrambi i ruoli, copre il 100%
                    if persona['tipo_persona'] == 'proprietario_inquilino':
                        percentuale = 100
                    else:
                        ruoli_presenti = unita_ruoli.get(persona['unita_id'], set())
                        if 'proprietario' in ruoli_presenti and 'inquilino' in ruoli_presenti:
                            percentuale = 50
                        else:
                            percentuale = 100
                else:  # personalizzato
                    percentuale = spesa['percentuale_proprietario'] if persona['tipo_persona'] in ['proprietario', 'proprietario_inquilino'] else spesa['percentuale_inquilino']

                # Calcola importo dovuto per questa spesa
                importo_dovuto = (spesa['importo'] * millesimi * percentuale / 100) / 1000
                importo_dovuto = round(importo_dovuto, 2)

                # Aggiungi alle spese per tabella
                tabella = spesa['tabella_millesimi']
                if tabella not in persona_data['spese_per_tabella']:
                    persona_data['spese_per_tabella'][tabella] = {
                        'tabella': tabella,
                        'totale_tabella': 0.0,
                        'spese': []
                    }

                persona_data['spese_per_tabella'][tabella]['totale_tabella'] += importo_dovuto
                persona_data['spese_per_tabella'][tabella]['spese'].append({
                    'spesa_id': spesa['id'],
                    'data_spesa': spesa['data_formatted'],
                    'importo_totale': spesa['importo'],
                    'importo_dovuto': importo_dovuto,
                    'logica_pi': spesa['logica_pi'],
                    'percentuale_proprietario': spesa['percentuale_proprietario'],
                    'percentuale_inquilino': spesa['percentuale_inquilino'],
                    'millesimi': millesimi,
                    'descrizione': spesa['descrizione']
                })

                persona_data['totale_dovuto'] += importo_dovuto

            # Converti le spese per tabella da dizionario a lista
            persona_data['spese_per_tabella'] = list(persona_data['spese_per_tabella'].values())

            # Arrotonda il totale
            persona_data['totale_dovuto'] = round(persona_data['totale_dovuto'], 2)

            # Arrotonda i totali per tabella e gli importi delle singole spese
            for tabella_data in persona_data['spese_per_tabella']:
                tabella_data['totale_tabella'] = round(tabella_data.get('totale_tabella', 0) or 0, 2)
                for sd in tabella_data['spese']:
                    sd['importo_dovuto'] = round(float(sd.get('importo_dovuto', 0) or 0), 2)

            result.append(persona_data)
            totale_generale += persona_data['totale_dovuto']

        conn.close()

        return jsonify({
            'ripartizione_dettagliata': result,
            'totale_generale': totale_generale,
            'tabella_filter': tabella_filter
        }), 200

    except Exception as e:
        log_error(str(e), f'get_ripartizione_dettagliata {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>/ripartizione/persona/<int:persona_id>', methods=['GET'])
@token_required
def get_ripartizione_persona(condo_id, persona_id):
    """Dettagli ripartizione per persona"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Verifica persona appartiene al condominio
        persona = Persona.find_by_id(persona_id)
        if not persona or persona.condominio_id != condo_id:
            return jsonify({'message': 'Persona non trovata'}), 404

        # Calcola ripartizione
        calculate_ripartizione_completa(condo_id)

        # Ottieni dettagli ripartizione per persona
        from database_universal import get_db
        conn = get_db()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT rs.*, s.descrizione, s.importo as importo_totale, s.tabella_millesimi
            FROM ripartizione_spese rs
            JOIN spese s ON rs.spesa_id = s.id
            WHERE rs.persona_id = ? AND rs.condominio_id = ?
            ORDER BY s.created_at
        """, (persona_id, condo_id))

        dettagli = cursor.fetchall()
        conn.close()

        result = []
        totale_dovuto = 0

        for dettaglio in dettagli:
            result.append({
                'spesa_id': dettaglio['spesa_id'],
                'descrizione': dettaglio['descrizione'],
                'importo_totale': dettaglio['importo_totale'],
                'tabella_millesimi': dettaglio['tabella_millesimi'],
                'importo_dovuto': dettaglio['importo_dovuto'],
                'anno': dettaglio['anno']
            })
            totale_dovuto += dettaglio['importo_dovuto']

        return jsonify({
            'persona': {
                'id': persona.id,
                'nome': persona.nome,
                'cognome': persona.cognome,
                'tipo_persona': persona.tipo_persona
            },
            'dettagli': result,
            'totale_dovuto': totale_dovuto
        }), 200

    except Exception as e:
        log_error(str(e), f'get_ripartizione_persona {condo_id} {persona_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>/ripartizione/recalcola', methods=['POST'])
@token_required
def recalcola_ripartizione(condo_id):
    """Forza recalcolo ripartizione"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Calcola ripartizione completa
        ripartizione = calculate_ripartizione_completa(condo_id)

        return jsonify({
            'message': 'Ripartizione ricalcolata con successo',
            'totale': sum(ripartizione.values())
        }), 200

    except Exception as e:
        log_error(str(e), f'recalcola_ripartizione {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

# ======================
# PREVENTIVO ANNUALE ENDPOINTS
# ======================

@app.route('/api/condominii/<int:condo_id>/preventivi', methods=['GET'])
@token_required
def get_preventivi(condo_id):
    """Lista preventivi per anni"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        preventivi = PreventivoAnnuale.get_by_condominio(condo_id)

        result = []
        for preventivo in preventivi:
            result.append({
                'id': preventivo.id,
                'anno': preventivo.anno,
                'importo_totale_preventivato': preventivo.importo_totale_preventivato,
                'importo_totale_speso': preventivo.importo_totale_speso,
                'differenza': preventivo.differenza,
                'note': preventivo.note,
                'created_at': preventivo.created_at,
                'updated_at': preventivo.updated_at
            })

        return jsonify(result), 200

    except Exception as e:
        log_error(str(e), f'get_preventivi {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>/preventivi/<int:anno>/genera', methods=['POST'])
@token_required
def genera_preventivo(condo_id, anno):
    """Genera preventivo anno"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Genera preventivo
        result = generate_preventivo_anno(condo_id, anno)

        return jsonify({
            'message': 'Preventivo generato con successo',
            'preventivo': {
                'id': result['preventivo'].id,
                'anno': result['preventivo'].anno,
                'importo_totale_preventivato': result['preventivo'].importo_totale_preventivato,
                'importo_totale_speso': result['preventivo'].importo_totale_speso,
                'note': result['preventivo'].note
            },
            'dettaglio': result['dettaglio']
        }), 201

    except Exception as e:
        log_error(str(e), f'genera_preventivo {condo_id} {anno}')
        return jsonify({'message': 'Errore del server'}), 500

# ======================
# SPESE PREVENTIVATE ENDPOINTS
# ======================

@app.route('/api/condominii/<int:condo_id>/spese-preventivate/<int:anno>', methods=['GET'])
@token_required
def get_spese_preventivate(condo_id, anno):
    """Lista spese preventivate per condominio e anno"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        spese = SpesaPreventivata.get_by_condominio_anno(condo_id, anno)

        result = []
        for spesa in spese:
            result.append({
                'id': spesa.id,
                'preventivo_id': spesa.preventivo_id,
                'descrizione': spesa.descrizione,
                'importo_previsto': spesa.importo_previsto,
                'tabella_millesimi': spesa.tabella_millesimi,
                'logica_pi': spesa.logica_pi,
                'percentuale_proprietario': spesa.percentuale_proprietario,
                'percentuale_inquilino': spesa.percentuale_inquilino,
                'mese_previsto': spesa.mese_previsto,
                'data_prevista': spesa.data_prevista,
                'note': spesa.note,
                'created_at': spesa.created_at
            })

        return jsonify(result), 200

    except Exception as e:
        log_error(str(e), f'get_spese_preventivate {condo_id} {anno}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/condominii/<int:condo_id>/spese-preventivate', methods=['POST'])
@token_required
def create_spesa_preventivata(condo_id):
    """Crea spesa preventivata"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        # Validazione base (riusa validate_spesa_data con adattamenti)
        errors = validate_spesa_data(data)
        if errors:
            return jsonify({'message': 'Dati non validi', 'errors': errors}), 400

        # Trova o crea preventivo per l'anno
        anno = data.get('anno', datetime.now().year + 1)
        preventivi = PreventivoAnnuale.get_by_condominio(condo_id)
        preventivo = None

        for p in preventivi:
            if p.anno == anno:
                preventivo = p
                break

        if not preventivo:
            preventivo = PreventivoAnnuale(
                condominio_id=condo_id,
                anno=anno,
                importo_totale_preventivato=0,
                note=f"Preventivo {anno}"
            )
            preventivo.save()

        # Crea spesa preventivata
        spesa = SpesaPreventivata(
            condominio_id=condo_id,
            preventivo_id=preventivo.id,
            descrizione=data['descrizione'],
            importo_previsto=data['importo'],
            tabella_millesimi=data['tabella_millesimi'],
            logica_pi=data['logica_pi'],
            percentuale_proprietario=data.get('percentuale_proprietario', 100),
            percentuale_inquilino=data.get('percentuale_inquilino', 0),
            mese_previsto=data.get('mese_previsto'),
            data_prevista=data.get('data_prevista'),
            note=data.get('note', '')
        )
        spesa.save()

        # Aggiorna totale preventivo
        spese_preventivate = SpesaPreventivata.get_by_preventivo(preventivo.id)
        nuovo_totale = sum(spesa.importo_previsto for spesa in spese_preventivate)
        preventivo.importo_totale_preventivato = nuovo_totale
        preventivo.save()

        return jsonify({
            'message': 'Spesa preventivata creata con successo',
            'spesa': {
                'id': spesa.id,
                'preventivo_id': spesa.preventivo_id,
                'descrizione': spesa.descrizione,
                'importo_previsto': spesa.importo_previsto,
                'tabella_millesimi': spesa.tabella_millesimi,
                'logica_pi': spesa.logica_pi,
                'percentuale_proprietario': spesa.percentuale_proprietario,
                'percentuale_inquilino': spesa.percentuale_inquilino,
                'mese_previsto': spesa.mese_previsto,
                'data_prevista': spesa.data_prevista,
                'note': spesa.note
            }
        }), 201

    except Exception as e:
        log_error(str(e), f'create_spesa_preventivata {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/spese-preventivate/<int:spesa_id>', methods=['PUT'])
@token_required
def update_spesa_preventivata(spesa_id):
    """Modifica spesa preventivata"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'message': 'Dati mancanti'}), 400

        spesa = SpesaPreventivata.find_by_id(spesa_id)
        if not spesa:
            return jsonify({'message': 'Spesa preventivata non trovata'}), 404

        # Verifica autorizzazione
        condominio = Condominio.find_by_id(spesa.condominio_id)
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Aggiorna campi
        if 'descrizione' in data:
            spesa.descrizione = data['descrizione']
        if 'importo_previsto' in data:
            spesa.importo_previsto = data['importo_previsto']
        if 'tabella_millesimi' in data:
            spesa.tabella_millesimi = data['tabella_millesimi']
        if 'logica_pi' in data:
            spesa.logica_pi = data['logica_pi']
        if 'percentuale_proprietario' in data:
            spesa.percentuale_proprietario = data['percentuale_proprietario']
        if 'percentuale_inquilino' in data:
            spesa.percentuale_inquilino = data['percentuale_inquilino']
        if 'mese_previsto' in data:
            spesa.mese_previsto = data['mese_previsto']
        if 'data_prevista' in data:
            spesa.data_prevista = data['data_prevista']
        if 'note' in data:
            spesa.note = data['note']

        spesa.save()

        # Aggiorna totale preventivo
        spese_preventivate = SpesaPreventivata.get_by_preventivo(spesa.preventivo_id)
        preventivo = PreventivoAnnuale.find_by_id(spesa.preventivo_id)
        preventivo.importo_totale_preventivato = sum(spesa.importo_previsto for spesa in spese_preventivate)
        preventivo.save()

        return jsonify({
            'message': 'Spesa preventivata aggiornata con successo',
            'spesa': {
                'id': spesa.id,
                'descrizione': spesa.descrizione,
                'importo_previsto': spesa.importo_previsto,
                'tabella_millesimi': spesa.tabella_millesimi,
                'logica_pi': spesa.logica_pi,
                'percentuale_proprietario': spesa.percentuale_proprietario,
                'percentuale_inquilino': spesa.percentuale_inquilino,
                'mese_previsto': spesa.mese_previsto,
                'data_prevista': spesa.data_prevista,
                'note': spesa.note
            }
        }), 200

    except Exception as e:
        log_error(str(e), f'update_spesa_preventivata {spesa_id}')
        return jsonify({'message': 'Errore del server'}), 500

@app.route('/api/spese-preventivate/<int:spesa_id>', methods=['DELETE'])
@token_required
def delete_spesa_preventivata(spesa_id):
    """Elimina spesa preventivata"""
    try:
        spesa = SpesaPreventivata.find_by_id(spesa_id)
        if not spesa:
            return jsonify({'message': 'Spesa preventivata non trovata'}), 404

        # Verifica autorizzazione
        condominio = Condominio.find_by_id(spesa.condominio_id)
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        preventivo_id = spesa.preventivo_id
        spesa.delete()

        # Aggiorna totale preventivo
        spese_preventivate = SpesaPreventivata.get_by_preventivo(preventivo_id)
        preventivo = PreventivoAnnuale.find_by_id(preventivo_id)
        preventivo.importo_totale_preventivato = sum(spesa.importo_previsto for spesa in spese_preventivate)
        preventivo.save()

        return jsonify({'message': 'Spesa preventivata eliminata con successo'}), 200

    except Exception as e:
        log_error(str(e), f'delete_spesa_preventivata {spesa_id}')
        return jsonify({'message': 'Errore del server'}), 500

# ======================
# CALCOLO PREVENTIVO ENDPOINTS
# ======================

@app.route('/api/condominii/<int:condo_id>/calcolo-preventivo/<int:anno>', methods=['GET'])
@token_required
def get_calcolo_preventivo(condo_id, anno):
    """Calcola ripartizione basata su preventivo"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Calcola ripartizione preventivo
        calc = calculate_ripartizione_preventivo(condo_id, anno)

        # Enrichment/naming per coerenza frontend
        # Mappa persona_id -> numero_unita
        from database_universal import get_db
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute('''
            SELECT p.id as persona_id, ui.numero_unita
            FROM persone p
            JOIN unita_immobiliari ui ON p.unita_id = ui.id
            WHERE p.condominio_id = ?
        ''', (condo_id,))
        num_map = {row['persona_id']: row['numero_unita'] for row in cursor.fetchall()}
        conn.close()

        rip_ui = []
        for item in calc['ripartizione']:
            rip_ui.append({
                'persona_id': item['persona_id'],
                'nome': item['nome'],
                'cognome': item['cognome'],
                'tipo_persona': item['tipo_persona'],
                'numero_unita': num_map.get(item['persona_id']),
                # rinomina per il frontend
                'importo_dovuto': round(float(item.get('importo_previsto_dovuto', 0) or 0), 2)
            })

        return jsonify({
            'message': 'Calcolo preventivo completato',
            'ripartizione': rip_ui,
            # rinomina chiave per il frontend
            'totale': round(float(calc.get('totale_previsto', 0) or 0), 2)
        }), 200

    except Exception as e:
        log_error(str(e), f'get_calcolo_preventivo {condo_id} {anno}')
        return jsonify({'message': 'Errore del server'}), 500

# ======================
# IMPORT/EXPORT ENDPOINTS
# ======================

@app.route('/api/condominii/<int:condo_id>/export', methods=['POST'])
@token_required
def export_condominio(condo_id):
    """Esporta condominio come JSON"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Esporta dati
        export_json = export_condominio_json(condo_id)
        if not export_json:
            return jsonify({'message': 'Errore esportazione'}), 500

        return jsonify({
            'message': 'Export completato con successo',
            'data': export_json
        }), 200

    except Exception as e:
        log_error(str(e), f'export_condominio {condo_id}')
        return jsonify({'message': 'Errore del server'}), 500

# ======================
# STAMPA WORD ENDPOINTS
# ======================

@app.route('/api/condominii/<int:condo_id>/stampa/spese', methods=['GET'])
@token_required
def stampa_spese(condo_id):
    """Genera documento Word con elenco spese"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Filtro opzionale per tabella
        tabella_filter = request.args.get('tabella')
        if tabella_filter and tabella_filter not in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'L']:
            return jsonify({'message': 'Tabella non valida'}), 400

        # Crea documento Word
        doc = Document()

        # Intestazione
        title = doc.add_heading(f'ELENCO SPESE CONDOMINIO "{condominio.nome.upper()}"', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sottotitolo con filtro tabella
        if tabella_filter:
            subtitle = doc.add_paragraph(f'Filtro Tabella Millesimi: {tabella_filter}')
        else:
            subtitle = doc.add_paragraph('Tutte le tabelle millesimi')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spazio

        # Informazioni condominio
        info_para = doc.add_paragraph()
        info_para.add_run('Indirizzo: ').bold = True
        info_para.add_run(f'{condominio.indirizzo or "N/D"}')

        info_para2 = doc.add_paragraph()
        info_para2.add_run('Numero Unità: ').bold = True
        info_para2.add_run(f'{condominio.num_unita}')

        info_para3 = doc.add_paragraph()
        info_para3.add_run('Data emissione: ').bold = True
        info_para3.add_run(f'{datetime.now().strftime("%d/%m/%Y")}')

        doc.add_paragraph()  # Spazio

        # Ottieni spese con raggruppamento per tabella millesimi
        from database_universal import get_db
        conn = get_db()
        cursor = conn.cursor()

        if tabella_filter:
            cursor.execute("""
                SELECT s.*
                FROM spese s
                WHERE s.condominio_id = ? AND s.tabella_millesimi = ?
                ORDER BY s.data_spesa DESC, s.created_at DESC
            """, (condo_id, tabella_filter))
        else:
            cursor.execute("""
                SELECT s.*
                FROM spese s
                WHERE s.condominio_id = ?
                ORDER BY s.tabella_millesimi, s.data_spesa DESC, s.created_at DESC
            """, (condo_id,))

        spese_data = cursor.fetchall()
        conn.close()

        if not spese_data:
            doc.add_paragraph('Nessuna spesa trovata per i criteri selezionati.')
        else:
            # Raggruppa spese per tabella millesimi
            spese_per_tabella = {}
            totale_generale = 0

            for spesa_row in spese_data:
                tabella = spesa_row['tabella_millesimi']
                if tabella not in spese_per_tabella:
                    spese_per_tabella[tabella] = {
                        'tabella': tabella,
                        'totale_tabella': 0,
                        'spese': []
                    }

                # Gestione del formato data
                if hasattr(spesa_row['data_spesa'], 'strftime'):
                    data_formattata = spesa_row['data_spesa'].strftime('%d/%m/%Y')
                elif isinstance(spesa_row['data_spesa'], str):
                    data_formattata = spesa_row['data_spesa']
                else:
                    data_formattata = 'N/D'

                spesa_dettaglio = {
                    'id': spesa_row['id'],
                    'descrizione': spesa_row['descrizione'],
                    'importo': spesa_row['importo'],
                    'data_spesa': data_formattata,
                    'logica_pi': spesa_row['logica_pi'],
                    'percentuale_proprietario': spesa_row['percentuale_proprietario'],
                    'percentuale_inquilino': spesa_row['percentuale_inquilino'],
                    'created_at': spesa_row['created_at']
                }

                spese_per_tabella[tabella]['spese'].append(spesa_dettaglio)
                spese_per_tabella[tabella]['totale_tabella'] += spesa_row['importo']
                totale_generale += spesa_row['importo']

            # Riepilogo generale
            doc.add_paragraph()
            summary_heading = doc.add_heading('RIEPILOGO SPESE PER TABELLA MILLESIMI', level=1)
            summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Tabella riassuntiva per tabella
            summary_table = doc.add_table(rows=1, cols=3)
            summary_table.style = 'Table Grid'
            summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Intestazioni tabella riassuntiva
            hdr_cells = summary_table.rows[0].cells
            hdr_cells[0].text = 'Tabella Millesimi'
            hdr_cells[1].text = 'Numero Spese'
            hdr_cells[2].text = 'Totale Tabella'

            # Formatta intestazioni
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    paragraph.runs[0].bold = True

            for tabella, data in spese_per_tabella.items():
                row_cells = summary_table.add_row().cells
                row_cells[0].text = tabella
                row_cells[1].text = str(len(data['spese']))
                row_cells[2].text = f'€ {data["totale_tabella"]:.2f}'

            # Totale generale
            doc.add_paragraph()
            total_para = doc.add_paragraph()
            total_para.add_run('TOTALE GENERALE SPESE: ').bold = True
            total_para.add_run(f'€ {totale_generale:.2f}')
            total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Dettaglio spese per tabella
            for tabella, data in spese_per_tabella.items():
                doc.add_paragraph()
                doc.add_paragraph().add_run("=" * 80).bold = True
                doc.add_paragraph()

                tabella_heading = doc.add_heading(f'TABELLA MILLESIMI: {tabella}', level=1)
                tabella_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Informazioni tabella
                tabella_info = doc.add_paragraph()
                tabella_info.add_run('Numero spese: ').bold = True
                tabella_info.add_run(f'{len(data["spese"])}')

                tabella_info2 = doc.add_paragraph()
                tabella_info2.add_run('Totale tabella: ').bold = True
                tabella_info2.add_run(f'€ {data["totale_tabella"]:.2f}')

                doc.add_paragraph()

                # Tabella dettaglio spese
                detail_table = doc.add_table(rows=1, cols=6)
                detail_table.style = 'Table Grid'

                # Intestazioni tabella dettaglio
                detail_hdr_cells = detail_table.rows[0].cells
                detail_hdr_cells[0].text = 'Data'
                detail_hdr_cells[1].text = 'Descrizione'
                detail_hdr_cells[2].text = 'Importo'
                detail_hdr_cells[3].text = 'Logica P/I'
                detail_hdr_cells[4].text = '% Proprietario'
                detail_hdr_cells[5].text = '% Inquilino'

                # Formatta intestazioni
                for cell in detail_hdr_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.runs[0].bold = True

                for spesa in data['spese']:
                    detail_row_cells = detail_table.add_row().cells
                    detail_row_cells[0].text = spesa['data_spesa']
                    detail_row_cells[1].text = spesa['descrizione']
                    detail_row_cells[2].text = f'€ {spesa["importo"]:.2f}'
                    detail_row_cells[3].text = spesa['logica_pi'].replace('_', ' ').title()
                    detail_row_cells[4].text = f'{spesa["percentuale_proprietario"]}%'
                    detail_row_cells[5].text = f'{spesa["percentuale_inquilino"]}%'

                # Subtotale tabella
                doc.add_paragraph()
                subtotal_para = doc.add_paragraph()
                subtotal_para.add_run(f'Subtotale Tabella {tabella}: ').bold = True
                subtotal_para.add_run(f'€ {data["totale_tabella"]:.2f}')
                subtotal_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.add_paragraph()

            # Statistiche finali
            doc.add_paragraph()
            doc.add_paragraph().add_run("=" * 80).bold = True
            doc.add_paragraph()

            stats_heading = doc.add_heading('STATISTICHE FINALI', level=1)
            stats_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            stats_para = doc.add_paragraph()
            stats_para.add_run('Numero totale tabelle: ').bold = True
            stats_para.add_run(f'{len(spese_per_tabella)}')

            stats_para2 = doc.add_paragraph()
            stats_para2.add_run('Numero totale spese: ').bold = True
            stats_para2.add_run(f'{len(spese_data)}')

            stats_para3 = doc.add_paragraph()
            stats_para3.add_run('Importo medio spesa: ').bold = True
            stats_para3.add_run(f'€ {totale_generale/len(spese_data):.2f}' if len(spese_data) > 0 else '€ 0.00')

            # Note finali
            doc.add_paragraph()
            note_para = doc.add_paragraph()
            note_para.add_run('Note: ').bold = True
            if tabella_filter:
                note_para.add_run(f'Questo elenco include solo le spese assegnate alla tabella millesimi {tabella_filter}.')
            else:
                note_para.add_run('Questo elenco include tutte le spese del condominio suddivise per tabelle millesimi.')

        # Prepara response
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"spese_{condominio.nome.replace(' ', '_')}"
        if tabella_filter:
            filename += f"_tabella_{tabella_filter}"
        filename += f"_{datetime.now().strftime('%d%m%Y')}.docx"

        response = make_response(buffer.getvalue())
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        return response

    except Exception as e:
        log_error(str(e), f'stampa_spese {condo_id}')
        return jsonify({'message': 'Errore durante la generazione del documento'}), 500

@app.route('/api/condominii/<int:condo_id>/stampa/ripartizione', methods=['GET'])
@token_required
def stampa_ripartizione(condo_id):
    """Genera documento Word con calcolo ripartizione"""
    try:
        # Verifica proprietà condominio
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Filtro opzionale per tabella
        tabella_filter = request.args.get('tabella')
        if tabella_filter and tabella_filter not in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'L']:
            return jsonify({'message': 'Tabella non valida'}), 400

        # Crea documento Word
        doc = Document()

        # Intestazione
        title = doc.add_heading(f'CALCOLO RIPARTIZIONE SPESE CONDOMINIO "{condominio.nome.upper()}"', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sottotitolo con filtro tabella
        if tabella_filter:
            subtitle = doc.add_paragraph(f'Ripartizione Tabella Millesimi: {tabella_filter}')
        else:
            subtitle = doc.add_paragraph('Ripartizione Tutte le Tabelle Millesimi')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spazio

        # Informazioni condominio
        info_para = doc.add_paragraph()
        info_para.add_run('Indirizzo: ').bold = True
        info_para.add_run(f'{condominio.indirizzo or "N/D"}')

        info_para2 = doc.add_paragraph()
        info_para2.add_run('Numero Unità: ').bold = True
        info_para2.add_run(f'{condominio.num_unita}')

        info_para3 = doc.add_paragraph()
        info_para3.add_run('Data calcolo: ').bold = True
        info_para3.add_run(f'{datetime.now().strftime("%d/%m/%Y")}')

        doc.add_paragraph()  # Spazio

        # Ottieni dati dettagliati come nell'interfaccia web
        from database_universal import get_db
        conn = get_db()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT p.*, ui.numero_unita
            FROM persone p
            JOIN unita_immobiliari ui ON p.unita_id = ui.id
            WHERE p.condominio_id = ?
            ORDER BY ui.numero_unita, p.cognome, p.nome
        """, (condo_id,))

        persone_data = cursor.fetchall()
        conn.close()

        risultati = []
        totale_generale = 0

        for persona_row in persone_data:
            # Crea oggetto Persona con numero_unita
            persona = Persona(
                id=persona_row['id'],
                nome=persona_row['nome'],
                cognome=persona_row['cognome'],
                email=persona_row['email'],
                tipo_persona=persona_row['tipo_persona'],
                unita_id=persona_row['unita_id']
            )
            # Aggiungi numero_unita come attributo
            persona.numero_unita = persona_row['numero_unita']

            from database_universal import get_db
            conn = get_db()
            cursor = conn.cursor()

            # Ottieni spese per persona con filtro tabella
            if tabella_filter:
                cursor.execute("""
                    SELECT s.*, rs.importo_dovuto
                    FROM spese s
                    LEFT JOIN ripartizione_spese rs ON s.id = rs.spesa_id AND rs.persona_id = ?
                    WHERE s.condominio_id = ? AND s.tabella_millesimi = ?
                    ORDER BY s.created_at DESC
                """, (persona.id, condo_id, tabella_filter))
            else:
                cursor.execute("""
                    SELECT s.*, rs.importo_dovuto
                    FROM spese s
                    LEFT JOIN ripartizione_spese rs ON s.id = rs.spesa_id AND rs.persona_id = ?
                    WHERE s.condominio_id = ?
                    ORDER BY s.tabella_millesimi, s.created_at DESC
                """, (persona.id, condo_id))

            spese_persona = cursor.fetchall()
            conn.close()

            if spese_persona:
                # Raggruppa spese per tabella
                spese_per_tabella = {}
                totale_persona = 0

                for spesa in spese_persona:
                    # Includi solo le spese con importo_dovuto valido
                    if spesa['importo_dovuto'] is not None and spesa['importo_dovuto'] > 0:
                        tabella = spesa['tabella_millesimi']
                        if tabella not in spese_per_tabella:
                            spese_per_tabella[tabella] = {
                                'tabella': tabella,
                                'totale_tabella': 0,
                                'spese': []
                            }

                        spesa_dettaglio = {
                            'spesa_id': spesa['id'],
                            'descrizione': spesa['descrizione'],
                            'importo_totale': spesa['importo'],
                            'importo_dovuto': spesa['importo_dovuto'],
                            'data_spesa': spesa['data_spesa'],
                            'logica_pi': spesa['logica_pi']
                        }

                        spese_per_tabella[tabella]['spese'].append(spesa_dettaglio)
                        spese_per_tabella[tabella]['totale_tabella'] += spesa['importo_dovuto']
                        totale_persona += spesa['importo_dovuto']

                # Formatta il risultato per questa persona
                persona_result = {
                    'persona': persona,
                    'totale_dovuto': totale_persona,
                    'spese_per_tabella': list(spese_per_tabella.values())
                }

                risultati.append(persona_result)
                totale_generale += totale_persona

        if not risultati:
            doc.add_paragraph('Nessuna ripartizione trovata per i criteri selezionati.')
        else:
            # Riepilogo generale
            doc.add_paragraph()
            summary_heading = doc.add_heading('RIEPILOGO GENERALE', level=1)
            summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Tabella riassuntiva
            summary_table = doc.add_table(rows=1, cols=4)
            summary_table.style = 'Table Grid'
            summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Intestazioni tabella riassuntiva
            hdr_cells = summary_table.rows[0].cells
            hdr_cells[0].text = 'Unità'
            hdr_cells[1].text = 'Intestatario'
            hdr_cells[2].text = 'Tipo'
            hdr_cells[3].text = 'Totale Dovuto'

            # Formatta intestazioni
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    paragraph.runs[0].bold = True

            for risultato in risultati:
                row_cells = summary_table.add_row().cells
                persona = risultato['persona']
                row_cells[0].text = str(persona.numero_unita)
                row_cells[1].text = f"{persona.cognome} {persona.nome}"
                row_cells[2].text = persona.tipo_persona.replace('_', ' ').title()
                row_cells[3].text = f'€ {risultato["totale_dovuto"]:.2f}'

            # Totale generale
            doc.add_paragraph()
            total_para = doc.add_paragraph()
            total_para.add_run('TOTALE GENERALE: ').bold = True
            total_para.add_run(f'€ {totale_generale:.2f}')
            total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Dettaglio per persona
            for i, risultato in enumerate(risultati):
                doc.add_paragraph()
                doc.add_paragraph().add_run("=" * 60).bold = True
                doc.add_paragraph()

                persona = risultato['persona']
                persona_heading = doc.add_heading(f'UNITÀ {persona.numero_unita} - {persona.cognome} {persona.nome}', level=1)
                persona_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                info_para = doc.add_paragraph()
                info_para.add_run('Tipo: ').bold = True
                info_para.add_run(persona.tipo_persona.replace('_', ' ').title())

                info_para2 = doc.add_paragraph()
                info_para2.add_run('Totale Dovuto: ').bold = True
                info_para2.add_run(f'€ {risultato["totale_dovuto"]:.2f}')

                doc.add_paragraph()

                # Dettaglio spese per tabella
                for tabella_data in risultato['spese_per_tabella']:
                    doc.add_paragraph().add_run(f'Tabella Millesimi: {tabella_data["tabella"]}').bold = True

                    # Tabella spese per questa tabella
                    spese_table = doc.add_table(rows=1, cols=4)
                    spese_table.style = 'Table Grid'

                    # Intestazioni tabella spese
                    spese_hdr_cells = spese_table.rows[0].cells
                    spese_hdr_cells[0].text = 'Data'
                    spese_hdr_cells[1].text = 'Descrizione'
                    spese_hdr_cells[2].text = 'Importo Totale'
                    spese_hdr_cells[3].text = 'Importo Dovuto'

                    # Formatta intestazioni
                    for cell in spese_hdr_cells:
                        for paragraph in cell.paragraphs:
                            paragraph.runs[0].bold = True

                    for spesa in tabella_data['spese']:
                        spese_row_cells = spese_table.add_row().cells

                        # Gestione del formato data
                        if hasattr(spesa['data_spesa'], 'strftime'):
                            data_formattata = spesa['data_spesa'].strftime('%d/%m/%Y')
                        elif isinstance(spesa['data_spesa'], str):
                            data_formattata = spesa['data_spesa']
                        else:
                            data_formattata = 'N/D'

                        spese_row_cells[0].text = data_formattata
                        spese_row_cells[1].text = spesa['descrizione']
                        spese_row_cells[2].text = f'€ {spesa["importo_totale"]:.2f}'
                        spese_row_cells[3].text = f'€ {spesa["importo_dovuto"]:.2f}'

                    # Subtotale tabella
                    doc.add_paragraph()
                    subtotal_para = doc.add_paragraph()
                    subtotal_para.add_run(f'Subtotale Tabella {tabella_data["tabella"]}: ').bold = True
                    subtotal_para.add_run(f'€ {tabella_data["totale_tabella"]:.2f}')
                    subtotal_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    doc.add_paragraph()

            # Note finali
            doc.add_paragraph()
            doc.add_paragraph().add_run("=" * 60).bold = True
            doc.add_paragraph()
            note_para = doc.add_paragraph()
            note_para.add_run('Note: ').bold = True
            if tabella_filter:
                note_para.add_run(f'Questa ripartizione include solo le spese assegnate alla tabella millesimi {tabella_filter}.')
            else:
                note_para.add_run('Questa ripartizione include tutte le spese del condominio suddivise per tabelle millesimi.')

        # Prepara response
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"ripartizione_{condominio.nome.replace(' ', '_')}"
        if tabella_filter:
            filename += f"_tabella_{tabella_filter}"
        filename += f"_{datetime.now().strftime('%d%m%Y')}.docx"

        response = make_response(buffer.getvalue())
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        return response

    except Exception as e:
        log_error(str(e), f'stampa_ripartizione {condo_id}')
        return jsonify({'message': 'Errore durante la generazione del documento'}), 500

@app.route('/api/condominii/<int:condo_id>/stampa/preventivo/<int:anno>', methods=['GET'])
@token_required
def stampa_preventivo(condo_id, anno):
    """Genera documento Word con preventivo annuale e ripartizione prevista"""
    try:
        # Verifica condominio e permessi
        condominio = Condominio.find_by_id(condo_id)
        if not condominio:
            return jsonify({'message': 'Condominio non trovato'}), 404
        if condominio.user_id != request.current_user_id:
            return jsonify({'message': 'Non autorizzato'}), 403

        # Dati preventivo
        tabella_filter = request.args.get('tabella')
        if tabella_filter and tabella_filter not in ['A','B','C','D','E','F','G','H','I','L']:
            return jsonify({'message': 'Tabella non valida'}), 400
        spese_prev = SpesaPreventivata.get_by_condominio_anno(condo_id, anno)
        if tabella_filter:
            spese_prev = [s for s in spese_prev if s.tabella_millesimi == tabella_filter]
        calc = calculate_ripartizione_preventivo(condo_id, anno, tabella_filter)

        # Crea documento
        doc = Document()

        # Intestazione
        title = doc.add_heading(f'PREVENTIVO ANNUALE {anno} - CONDOMINIO "{condominio.nome.upper()}"', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle_text = 'Ripartizione prevista su base millesimale'
        if tabella_filter:
            subtitle_text += f' — Tabella {tabella_filter}'
        subtitle = doc.add_paragraph(subtitle_text)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Info condominio
        info = doc.add_paragraph()
        info.add_run('Indirizzo: ').bold = True
        info.add_run(f'{condominio.indirizzo or "N/D"}')

        info2 = doc.add_paragraph()
        info2.add_run('Numero Unità: ').bold = True
        info2.add_run(f'{condominio.num_unita}')

        info3 = doc.add_paragraph()
        info3.add_run('Data documento: ').bold = True
        info3.add_run(datetime.now().strftime('%d/%m/%Y'))

        doc.add_paragraph()

        # Riepilogo importi
        totale_previsto = sum((s.importo_previsto or 0) for s in spese_prev)
        tot_para = doc.add_paragraph()
        tot_para.add_run('Totale Spese Preventivate: ').bold = True
        tot_para.add_run(f"€ {totale_previsto:.2f}")

        if calc and isinstance(calc, dict):
            tot_para2 = doc.add_paragraph()
            tot_para2.add_run('Totale da ripartire (calcolo): ').bold = True
            tot_para2.add_run(f"€ {float(calc.get('totale', 0)):.2f}")

        doc.add_paragraph()

        # Tabella spese preventivate
        if spese_prev:
            doc.add_paragraph().add_run('Elenco Spese Preventivate').bold = True
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Data/Mese'
            hdr[1].text = 'Descrizione'
            hdr[2].text = 'Tabella'
            hdr[3].text = 'Logica P/I'
            hdr[4].text = 'Importo Previsto'
            for cell in hdr:
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
            for s in spese_prev:
                row = table.add_row().cells
                if s.data_prevista:
                    data_disp = s.data_prevista
                elif s.mese_previsto:
                    data_disp = f"Mese {s.mese_previsto}"
                else:
                    data_disp = '-'
                row[0].text = str(data_disp)
                row[1].text = s.descrizione or ''
                row[2].text = s.tabella_millesimi or ''
                row[3].text = s.logica_pi or ''
                row[4].text = f"€ {float(s.importo_previsto or 0):.2f}"

            doc.add_paragraph()

        # Dettaglio ripartizione per persona
        if calc and isinstance(calc, dict) and calc.get('ripartizione'):
            doc.add_paragraph().add_run('Ripartizione per Persona').bold = True
            for p in calc['ripartizione']:
                blocco = doc.add_paragraph()
                blocco.add_run(f"Unità {p.get('numero_unita','-')} - {p.get('cognome','')} {p.get('nome','')}").bold = True
                doc.add_paragraph(f"Importo Dovuto: € {float(p.get('importo_dovuto',0)):.2f}")

        # Response
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"preventivo_{anno}_{condominio.nome.replace(' ', '_')}_{datetime.now().strftime('%d%m%Y')}.docx"
        response = make_response(buffer.getvalue())
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return response

    except Exception as e:
        log_error(str(e), f'stampa_preventivo {condo_id} {anno}')
        return jsonify({'message': 'Errore durante la generazione del documento'}), 500
# ======================
# ERROR HANDLERS
# ======================

@app.errorhandler(404)
def not_found(error):
    return jsonify({'message': 'Risorsa non trovata'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'message': 'Errore interno del server'}), 500

@app.errorhandler(405)
def method_not_allowed(error):
    return jsonify({'message': 'Metodo non consentito'}), 405

# ======================
# MAIN
# ======================

if __name__ == '__main__':
    # Inizializza database
    init_db()
    create_default_user()

    print("CONDOMINIO NUOVO - WebApp")
    print("Server in esecuzione su http://localhost:5000")
    print("Login di default: admin / admin123")
    print("-" * 50)

    app.run(debug=True, host='0.0.0.0', port=5000)
